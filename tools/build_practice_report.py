from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
DOCX_OUT = REPORTS / "practice_report.docx"
PDF_OUT = REPORTS / "practice_report.pdf"

STUDENT = "Карпов Александр Владимирович"
GROUP = "251-371"
CODE = "10.05.03"
PROFILE = "Информационная безопасность автоматизированных систем"
PROJECT = "Eltech"
REPO_URL = "https://github.com/AlexanderKarpov07/Project-practice-2026"


def set_doc_style(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 14), ("Heading 3", 14)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5


def add_center(doc, text, size=14, bold=False, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    return p


def add_plain(doc, text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = bold
    return p


def add_left_no_indent(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def build_docx():
    REPORTS.mkdir(exist_ok=True)
    doc = Document()
    set_doc_style(doc)

    add_center(doc, "Федеральное государственное автономное образовательное учреждение высшего образования")
    add_center(doc, "«МОСКОВСКИЙ ПОЛИТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»", bold=True)
    add_center(doc, "Факультет информационных технологий\nКафедра «информационная безопасность»")
    add_center(doc, f"Направление подготовки/ специальность: {CODE} {PROFILE}")
    for _ in range(3):
        doc.add_paragraph()
    add_center(doc, "ОТЧЕТ", size=16, bold=True)
    add_center(doc, "по проектной практике", size=16, bold=True)
    doc.add_paragraph()
    add_left_no_indent(doc, f"Студент: {STUDENT}    Группа: {GROUP}")
    add_left_no_indent(doc, "Место прохождения практики: Московский Политех, кафедра информационной безопасности.")
    add_left_no_indent(doc, "Отчет принят с оценкой _______________ Дата ________________________")
    add_left_no_indent(doc, "Руководитель практики: Александр Юрьевич Гнешев")
    for _ in range(5):
        doc.add_paragraph()
    add_center(doc, "Москва 2026")
    doc.add_page_break()

    add_center(doc, "ОГЛАВЛЕНИЕ", bold=True)
    contents = [
        "ВВЕДЕНИЕ",
        "1. Общая информация о проекте",
        "1.1. Название проекта",
        "1.2. Цели и задачи проекта",
        "2. Общая характеристика деятельности организации (заказчика проекта)",
        "2.1. Наименование заказчика",
        "2.2. Организационная структура",
        "2.3. Описание деятельности",
        "3. Описание задания по проектной практике",
        "4. Описание достигнутых результатов по проектной практике",
        "5. Онлайн-курсы по искусственному интеллекту (LLM)",
        "6. Взаимодействие с организацией-партнёром",
        "7. Вариативная часть",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ",
        "ПРИЛОЖЕНИЯ",
    ]
    for item in contents:
        add_left_no_indent(doc, item)
    doc.add_page_break()

    add_heading(doc, "ВВЕДЕНИЕ", 1)
    add_plain(
        doc,
        f"Настоящий отчет составлен по итогам прохождения проектной практики студентом группы {GROUP} направления "
        f"«{PROFILE}». В рамках практики выполнялась работа по оформлению материалов проекта «{PROJECT}», "
        "созданию проектной документации, разработке статического сайта и подготовке итоговых отчетных файлов."
    )
    add_plain(
        doc,
        "Проектная практика позволила закрепить навыки работы с репозиторием, Markdown-документацией, HTML/CSS, "
        "структурированием проектных материалов и подготовкой документации для публикации на GitHub."
    )
    add_plain(doc, f"Ссылка на репозиторий с выполненными заданиями практики GitHub URL: {REPO_URL}.")

    add_heading(doc, "1. Общая информация о проекте", 1)
    add_heading(doc, "1.1. Название проекта", 2)
    add_plain(doc, "Название проекта: Eltech.")
    add_plain(doc, "Тематика проекта: мобильная разработка, цифровые сервисы университета, Kotlin Multiplatform.")
    add_plain(doc, "Руководитель проекта: Станислав Андреевич Клепиков.")

    add_heading(doc, "1.2. Цели и задачи проекта", 2)
    add_plain(
        doc,
        "Основная цель проекта - создать полноценное кроссплатформенное мобильное приложение для удобного доступа "
        "к информации и сервисам Московского Политеха со смартфонов."
    )
    add_plain(
        doc,
        "Актуальность проекта связана с тем, что значительная часть университетских сервисов используется через "
        "веб-интерфейс личного кабинета, который не всегда удобен на мобильных устройствах. Мобильное приложение "
        "должно упростить доступ к расписанию, LMS/Moodle, учебным материалам и другим повседневным сервисам."
    )
    add_left_no_indent(doc, "Задачи проекта:", bold=True)
    add_bullet(doc, "перенести основной функционал личного кабинета в формат мобильного приложения;")
    add_bullet(doc, "разработать SDK_Eltech и общую бизнес-логику на базе Kotlin Multiplatform;")
    add_bullet(doc, "реализовать экраны расписания, авторизации, LMS и физической культуры;")
    add_bullet(doc, "подготовить архитектуру, пригодную для дальнейшего развития приложения на Android и iOS;")
    add_bullet(doc, "оформить интерфейс в визуальном стиле Московского Политеха.")

    add_heading(doc, "2. Общая характеристика деятельности организации (заказчика проекта)", 1)
    add_heading(doc, "2.1. Наименование заказчика", 2)
    add_plain(doc, "Организация-партнер и основная площадка реализации проекта - Московский Политех.")
    add_plain(doc, "Проект связан с цифровой образовательной средой университета и ориентирован на студентов и преподавателей.")

    add_heading(doc, "2.2. Организационная структура", 2)
    add_plain(
        doc,
        "Московский Политех является образовательной организацией высшего образования. В рамках проекта взаимодействие "
        "строится вокруг цифровых сервисов университета: личного кабинета, расписания, LMS/Moodle и связанных информационных систем."
    )
    add_plain(
        doc,
        "Команда проекта распределяет задачи по направлениям: SDK, доменный слой, кэширование, пользовательский интерфейс, "
        "отдельные экраны приложения и документация."
    )

    add_heading(doc, "2.3. Описание деятельности", 2)
    add_plain(
        doc,
        "Деятельность организации связана с обучением студентов, сопровождением образовательного процесса и развитием "
        "цифровой инфраструктуры. Проект Eltech поддерживает эти направления за счет создания удобного мобильного клиента, "
        "который объединяет учебные сервисы в одном приложении."
    )

    add_heading(doc, "3. Описание задания по проектной практике", 1)
    add_plain(
        doc,
        "В соответствии с заданием по проектной практике необходимо было выполнить обязательную часть: настроить работу с "
        "репозиторием, изучить Markdown, подготовить документацию проекта, создать статический сайт и оформить итоговый отчет."
    )
    add_plain(doc, "В ходе работы были выполнены следующие этапы:")
    add_bullet(doc, "создана структура репозитория с папками docs, site, task, src и reports;")
    add_bullet(doc, "подготовлен файл README.md с информацией о студенте, проекте и структуре репозитория;")
    add_bullet(doc, "создана Markdown-документация проекта в папке docs;")
    add_bullet(doc, "разработан статический сайт проекта Eltech в папке site;")
    add_bullet(doc, "изучены онлайн-курсы и материалы по большим языковым моделям, подготовлен docs/llm_report.md;")
    add_bullet(doc, "подготовлен отчёт о взаимодействии с организацией-партнёром docs/partner_report.md;")
    add_bullet(doc, "подготовлено техническое руководство по вариативной части docs/variant_guide.md;")
    add_bullet(doc, "подготовлены итоговые файлы отчета в форматах DOCX и PDF.")

    add_heading(doc, "4. Описание достигнутых результатов по проектной практике", 1)
    add_plain(
        doc,
        "В папке docs подготовлены четыре обязательных Markdown-документа: project.md, members.md, journal.md и resources.md. "
        "Документы содержат описание проекта, цели и задачи, список участников, журнал прогресса и полезные ресурсы. "
        "Дополнительно подготовлен файл llm_report.md с отчётом по онлайн-курсам LLM."
    )
    add_plain(
        doc,
        "В папке site разработан статический сайт. Он включает главную страницу, страницу с описанием проекта, страницу участников, "
        "журнал прогресса и страницу ресурсов. В сайт добавлены локальные графические материалы: макет мобильного приложения и схема архитектуры."
    )
    add_plain(
        doc,
        "Сайт открывается в браузере без серверных настроек через файл site/index.html. Навигация между страницами реализована обычными HTML-ссылками."
    )
    add_plain(
        doc,
        "Также были подготовлены дополнительные отчётные материалы: docs/partner_report.md, docs/variant_guide.md и памятка SUBMISSION.md, "
        "в которой описан порядок создания задач, коммитов и загрузки материалов в репозиторий."
    )

    add_heading(doc, "5. Онлайн-курсы по искусственному интеллекту (LLM)", 1)
    add_plain(
        doc,
        "В рамках пункта 1.4 задания были изучены онлайн-курсы и учебные материалы, посвящённые большим языковым моделям, "
        "промпт-инжинирингу, этическим аспектам применения ИИ и использованию LLM в разработке и документировании."
    )
    add_left_no_indent(doc, "Курс 1: «Введение в LLM и архитектуру Transformer» (Hugging Face).", bold=True)
    add_plain(
        doc,
        "Изучены базовые принципы работы больших языковых моделей: архитектура Transformer, механизм внимания, токенизация, "
        "жизненный цикл модели и задачи обработки естественного языка."
    )
    add_left_no_indent(doc, "Курс 2: «Промпт-инжиниринг» (материалы по работе с GigaChat).", bold=True)
    add_plain(
        doc,
        "Освоены приёмы постановки роли и контекста, уточнения формата ответа, декомпозиции сложной задачи и проверки результата."
    )
    add_left_no_indent(doc, "Курс 3: «Этические аспекты использования ИИ» (Habr).", bold=True)
    add_plain(
        doc,
        "Рассмотрены риски недостоверных ответов, вопросы защиты персональных данных, необходимость проверки фактов и прозрачного использования ИИ-инструментов."
    )
    add_left_no_indent(doc, "Курс 4: «Применение LLM в разработке и документировании» (Hexlet).", bold=True)
    add_plain(
        doc,
        "Изучены способы применения LLM для подготовки README, технической документации, структуры сайта, объяснения кода и проверки полноты требований."
    )
    add_plain(
        doc,
        "По итогам изучения подготовлен файл docs/llm_report.md. Полученные знания применялись при создании Markdown-документации, "
        "структурировании страниц сайта и подготовке итогового отчёта по практике."
    )

    add_heading(doc, "6. Взаимодействие с организацией-партнёром", 1)
    add_plain(
        doc,
        "В рамках пункта 1.5 задания было описано взаимодействие с организацией-партнёром проекта. Для проекта Eltech такой организацией "
        "является Московский Политех, так как приложение связано с цифровой образовательной средой университета."
    )
    add_plain(
        doc,
        "В отчёте docs/partner_report.md рассмотрены потребности студентов, особенности мобильных сценариев, связь проекта с личным кабинетом, "
        "расписанием, LMS/Moodle и экраном физической культуры."
    )
    add_plain(
        doc,
        "Для моего направления это взаимодействие важно тем, что экран физической культуры должен отображать реальные данные студента: "
        "баллы, посещения и статус выполнения требований."
    )

    add_heading(doc, "7. Вариативная часть", 1)
    add_plain(
        doc,
        "Вариативная часть представлена как вклад в открытый проект Eltech. В материалах проекта отражен личный вклад Карпова Александра "
        "в разработку экрана физической культуры."
    )
    add_plain(
        doc,
        "Экран физической культуры предназначен для отображения прогресса студента по физкультуре: накопленных баллов, посещений занятий "
        "и информации, получаемой из личного кабинета. Этот функционал является частью общего пользовательского сценария мобильного приложения."
    )
    add_plain(
        doc,
        "В отчётных материалах вклад зафиксирован в файле docs/members.md и в журнале прогресса docs/journal.md. "
        "Дополнительно подготовлено техническое руководство docs/variant_guide.md, описывающее архитектуру и логику экрана физической культуры. В папке src размещены исходные Kotlin-файлы: PhysicalScreen.kt, GroupJournalScreen.kt и HomeScreen.kt."
    )

    add_heading(doc, "ЗАКЛЮЧЕНИЕ", 1)
    add_plain(
        doc,
        "В ходе проектной практики были выполнены основные требования задания: подготовлена документация проекта, создан сайт, "
        "оформлена структура репозитория и сформирован итоговый отчет. Работа позволила закрепить навыки оформления проектных материалов "
        "и публикации результатов в системе контроля версий."
    )
    add_plain(
        doc,
        "Практическая ценность выполненных задач заключается в том, что материалы проекта Eltech стали более структурированными и удобными "
        "для просмотра. Документация и сайт могут использоваться для представления проекта, знакомства с его целями, командой и ходом разработки."
    )

    add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ", 1)
    sources = [
        "Официальный сайт Московского Политеха. URL: https://mospolytech.ru/",
        "Документация Kotlin Multiplatform. URL: https://kotlinlang.org/docs/multiplatform.html",
        "JetBrains Kotlin Multiplatform. URL: https://www.jetbrains.com/kotlin-multiplatform/",
        "Android Developers Documentation. URL: https://developer.android.com/",
        "Material Design 3. URL: https://m3.material.io/",
        "Pro Git Book. URL: https://git-scm.com/book/ru/v2",
        "GitHub Docs. URL: https://docs.github.com/",
    ]
    for src in sources:
        add_left_no_indent(doc, src)

    add_heading(doc, "ПРИЛОЖЕНИЯ", 1)
    add_left_no_indent(doc, "Приложение А. Структура репозитория", bold=True)
    for line in [
        "README.md - общая информация о практике и проекте.",
        "docs/project.md - описание проекта Eltech.",
        "docs/members.md - участники и личный вклад.",
        "docs/journal.md - журнал прогресса.",
        "docs/resources.md - полезные ресурсы.",
        "docs/llm_report.md - отчёт по онлайн-курсам LLM.",
        "docs/partner_report.md - отчёт о взаимодействии с организацией-партнёром.",
        "docs/variant_guide.md - техническое руководство по экрану физической культуры.",
        "site/index.html - главная страница сайта.",
        "site/about.html - страница «О проекте».",
        "site/members.html - страница «Участники».",
        "site/journal.html - страница «Журнал».",
        "site/resources.html - страница «Ресурсы».",
        "reports/practice_report.docx - итоговый отчет в формате DOCX.",
        "reports/practice_report.pdf - итоговый отчет в формате PDF.",
    ]:
        add_bullet(doc, line)

    add_left_no_indent(doc, "Приложение Б. Краткое описание личного вклада", bold=True)
    add_plain(
        doc,
        "Карпов Александр Владимирович участвовал в направлении разработки экрана физической культуры приложения Eltech. "
        "В рамках отчётных материалов этот вклад отражён в документации участников и журнале проекта."
    )

    doc.save(DOCX_OUT)


def pdf_styles():
    pdfmetrics.registerFont(TTFont("TimesNewRoman", "C:/Windows/Fonts/times.ttf"))
    pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", "C:/Windows/Fonts/timesbd.ttf"))
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("TitleRu", fontName="TimesNewRoman-Bold", fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=10))
    styles.add(ParagraphStyle("CenterRu", fontName="TimesNewRoman", fontSize=12, leading=16, alignment=TA_CENTER, spaceAfter=4))
    styles.add(ParagraphStyle("H1Ru", fontName="TimesNewRoman-Bold", fontSize=14, leading=18, alignment=TA_LEFT, textColor="black", spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle("BodyRu", fontName="TimesNewRoman", fontSize=12, leading=17, alignment=TA_LEFT, firstLineIndent=1.25 * cm, spaceAfter=4))
    styles.add(ParagraphStyle("NoIndentRu", fontName="TimesNewRoman", fontSize=12, leading=17, alignment=TA_LEFT, spaceAfter=3))
    return styles


def safe(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def para(text, style):
    return Paragraph(safe(text), style)


def build_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=3 * cm,
        rightMargin=1.5 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="Итоговый отчет по проектной практике",
        author=STUDENT,
    )
    story = []
    for text in [
        "Федеральное государственное автономное образовательное учреждение высшего образования",
        "«МОСКОВСКИЙ ПОЛИТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»",
        "Факультет информационных технологий",
        "Кафедра «информационная безопасность»",
        f"Направление подготовки/ специальность: {CODE} {PROFILE}",
    ]:
        story.append(para(text, styles["CenterRu"]))
    story += [Spacer(1, 42), para("ОТЧЕТ", styles["TitleRu"]), para("по проектной практике", styles["TitleRu"]), Spacer(1, 24)]
    for text in [
        f"Студент: {STUDENT}    Группа: {GROUP}",
        "Место прохождения практики: Московский Политех, кафедра информационной безопасности.",
        "Отчет принят с оценкой _______________ Дата ________________________",
        "Руководитель практики: Александр Юрьевич Гнешев",
    ]:
        story.append(para(text, styles["NoIndentRu"]))
    story += [Spacer(1, 120), para("Москва 2026", styles["CenterRu"]), PageBreak()]

    story.append(para("ОГЛАВЛЕНИЕ", styles["TitleRu"]))
    for item in [
        "ВВЕДЕНИЕ",
        "1. Общая информация о проекте",
        "2. Общая характеристика деятельности организации",
        "3. Описание задания по проектной практике",
        "4. Описание достигнутых результатов",
        "5. Онлайн-курсы по искусственному интеллекту (LLM)",
        "6. Взаимодействие с организацией-партнёром",
        "7. Вариативная часть",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ",
        "ПРИЛОЖЕНИЯ",
    ]:
        story.append(para(item, styles["NoIndentRu"]))
    story.append(PageBreak())

    blocks = [
        ("ВВЕДЕНИЕ", [
            f"Настоящий отчет составлен по итогам прохождения проектной практики студентом группы {GROUP} направления «{PROFILE}».",
            "В рамках практики выполнялась работа по оформлению материалов проекта Eltech, созданию документации, разработке статического сайта и подготовке итоговых отчетных файлов.",
            f"Ссылка на репозиторий с выполненными заданиями практики GitHub URL: {REPO_URL}.",
        ]),
        ("1. Общая информация о проекте", [
            "Название проекта: Eltech.",
            "Тематика проекта: мобильная разработка, цифровые сервисы университета, Kotlin Multiplatform.",
            "Основная цель проекта - создать кроссплатформенное мобильное приложение для удобного доступа к информации и сервисам Московского Политеха.",
            "Задачи проекта включают перенос функций личного кабинета, разработку SDK_Eltech, интеграцию с LMS/Moodle, создание экранов расписания, авторизации и физической культуры.",
        ]),
        ("2. Общая характеристика деятельности организации", [
            "Организация-партнер и основная площадка реализации проекта - Московский Политех.",
            "Проект связан с цифровой образовательной средой университета и ориентирован на студентов и преподавателей.",
            "Eltech поддерживает развитие цифровой инфраструктуры университета за счет создания удобного мобильного клиента.",
        ]),
        ("3. Описание задания по проектной практике", [
            "В соответствии с заданием необходимо было настроить работу с репозиторием, изучить Markdown, подготовить документацию проекта, создать статический сайт и оформить итоговый отчет.",
            "В репозитории подготовлены папки docs, site, task, src и reports, а также корневой README.md.",
        ]),
        ("4. Описание достигнутых результатов по проектной практике", [
            "В папке docs созданы файлы project.md, members.md, journal.md и resources.md.",
            "Дополнительно подготовлен файл docs/llm_report.md с отчётом по онлайн-курсам LLM.",
            "Также подготовлены docs/partner_report.md и docs/variant_guide.md.",
            "В папке site разработан статический сайт с главной страницей, описанием проекта, участниками, журналом и ресурсами.",
            "Сайт содержит локальные графические материалы и открывается без серверных настроек через site/index.html.",
        ]),
        ("5. Онлайн-курсы по искусственному интеллекту (LLM)", [
            "В рамках пункта 1.4 были изучены четыре учебных направления: введение в LLM и Transformer, промпт-инжиниринг, этические аспекты ИИ и применение LLM в разработке и документировании.",
            "Курс 1: материалы Hugging Face по архитектуре Transformer, механизму внимания и токенизации.",
            "Курс 2: материалы по промпт-инжинирингу и работе с GigaChat.",
            "Курс 3: статьи Habr об этических аспектах использования искусственного интеллекта.",
            "Курс 4: материалы Hexlet о применении LLM в разработке и подготовке документации.",
            "По итогам подготовлен файл docs/llm_report.md.",
        ]),
        ("6. Взаимодействие с организацией-партнёром", [
            "Для проекта Eltech организацией-партнёром является Московский Политех, так как приложение связано с его цифровой образовательной средой.",
            "В docs/partner_report.md описаны потребности студентов, мобильные сценарии и связь проекта с личным кабинетом, расписанием, LMS/Moodle и экраном физической культуры.",
        ]),
        ("7. Вариативная часть", [
            "Вариативная часть представлена как вклад в открытый проект Eltech.",
            "В материалах проекта отражен личный вклад Карпова Александра в разработку экрана физической культуры.",
            "Экран предназначен для отображения баллов, посещений занятий и связанных данных из личного кабинета.",
            "В папке src размещены исходные Kotlin-файлы PhysicalScreen.kt, GroupJournalScreen.kt и HomeScreen.kt. Дополнительно подготовлено техническое руководство docs/variant_guide.md.",
        ]),
        ("ЗАКЛЮЧЕНИЕ", [
            "В ходе проектной практики были выполнены основные требования задания: подготовлена документация проекта, создан сайт, оформлена структура репозитория и сформирован итоговый отчет.",
            "Практическая ценность работы заключается в структурировании материалов проекта Eltech и подготовке их к публикации и проверке.",
        ]),
        ("СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ", [
            "Официальный сайт Московского Политеха. URL: https://mospolytech.ru/",
            "Документация Kotlin Multiplatform. URL: https://kotlinlang.org/docs/multiplatform.html",
            "Android Developers Documentation. URL: https://developer.android.com/",
            "GitHub Docs. URL: https://docs.github.com/",
        ]),
        ("ПРИЛОЖЕНИЯ", [
            "Приложение А. Структура репозитория: README.md, docs/, site/, task/, src/, reports/.",
            "Приложение Б. Краткое описание личного вклада: участие в направлении разработки экрана физической культуры приложения Eltech.",
        ]),
    ]
    for title, paragraphs in blocks:
        story.append(para(title, styles["H1Ru"]))
        for text in paragraphs:
            story.append(para(text, styles["BodyRu"]))

    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_OUT)
    print(PDF_OUT)
